"""
CV Parser Module - Extracts text from PDF and DOCX files
"""

import pdfplumber
import io
from docx import Document


def parse_pdf(file_bytes: bytes) -> dict:
    """
    Parse PDF file and extract text.
    Returns dict with: text, is_readable, page_count, has_images, has_tables
    """
    result = {
        "text": "",
        "is_readable": False,
        "page_count": 0,
        "has_images": False,
        "has_tables": False,
        "error": None,
    }

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            result["page_count"] = len(pdf.pages)
            all_text = []

            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    all_text.append(page_text)
                    result["is_readable"] = True

                # Check for tables
                tables = page.extract_tables()
                if tables:
                    result["has_tables"] = True

                # Check for images
                if page.images and len(page.images) > 0:
                    result["has_images"] = True

            result["text"] = "\n".join(all_text)

            # If no text extracted at all, likely image-based PDF
            if not result["text"].strip():
                result["is_readable"] = False
                result["error"] = (
                    "PDF appears to be image-based (scanned). ATS cannot read it."
                )

    except Exception as e:
        result["error"] = f"Error parsing PDF: {str(e)}"
        result["is_readable"] = False

    return result


def parse_docx(file_bytes: bytes) -> dict:
    """
    Parse DOCX file and extract text.
    Returns dict with: text, is_readable, has_tables, has_images
    """
    result = {
        "text": "",
        "is_readable": False,
        "page_count": 1,
        "has_images": False,
        "has_tables": False,
        "error": None,
    }

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        result["text"] = "\n".join(paragraphs)
        result["is_readable"] = len(result["text"].strip()) > 0

        # Check for tables
        if doc.tables and len(doc.tables) > 0:
            result["has_tables"] = True

        # Check for images in relationships
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                result["has_images"] = True
                break

        if not result["is_readable"]:
            result["error"] = "DOCX file appears to be empty or contains only images."

    except Exception as e:
        result["error"] = f"Error parsing DOCX: {str(e)}"
        result["is_readable"] = False

    return result


def parse_cv(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point - parses CV based on file extension.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return parse_docx(file_bytes)
    else:
        return {
            "text": "",
            "is_readable": False,
            "page_count": 0,
            "has_images": False,
            "has_tables": False,
            "error": f"Unsupported file format: {filename}. Please upload PDF or DOCX.",
        }
